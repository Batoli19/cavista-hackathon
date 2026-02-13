from pathlib import Path
from datetime import date
from typing import Dict, Any
import subprocess

def export_plan_to_word(project: Dict[str, Any]) -> str:
    from docx import Document  # python-docx
    out_dir = Path(__file__).resolve().parent.parent / "data"
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f'{project["name"].replace(" ", "_")}_plan_{date.today().isoformat()}.docx'
    out_path = out_dir / filename

    doc = Document()
    doc.add_heading(f'Project Plan: {project["name"]}', level=1)
    if project.get("deadline"):
        doc.add_paragraph(f'Deadline: {project["deadline"]}')
    if project.get("description"):
        doc.add_paragraph(project["description"])

    doc.add_heading("Tasks", level=2)
    for t in project.get("tasks", []):
        line = f'- {t.get("id")} | {t.get("name")} | {t.get("duration_days")} day(s)'
        if t.get("depends_on"):
            line += f' | depends on: {", ".join(t["depends_on"])}'
        doc.add_paragraph(line)

    doc.save(str(out_path))

    # Open file
    subprocess.Popen(["cmd", "/c", "start", "", str(out_path)], shell=True)
    return str(out_path)

def export_schedule_to_excel(project: Dict[str, Any], schedule: list[Dict[str, Any]]) -> str:
    from openpyxl import Workbook  # openpyxl
    out_dir = Path(__file__).resolve().parent.parent / "data"
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f'{project["name"].replace(" ", "_")}_schedule_{date.today().isoformat()}.xlsx'
    out_path = out_dir / filename

    wb = Workbook()
    ws = wb.active
    ws.title = "Schedule"
    ws.append(["Task ID", "Task Name", "Start", "End", "Duration Days", "Delay Days", "Status", "Depends On"])

    for t in schedule:
        ws.append([
            t.get("id"),
            t.get("name"),
            t.get("start"),
            t.get("end"),
            t.get("duration_days"),
            t.get("delay_days"),
            t.get("status"),
            ", ".join(t.get("depends_on", [])),
        ])

    wb.save(str(out_path))
    subprocess.Popen(["cmd", "/c", "start", "", str(out_path)], shell=True)
    return str(out_path)
